"""
Document extraction and preprocessing for requirement entry.

Used before intake / understanding: produces a single normalized text string from uploads
or passes through manual text after the same preprocessing when desired.
"""

from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass
from typing import Final

# --- Extraction ---


class DocumentExtractionError(Exception):
    """Raised when a file cannot be read or decoded."""


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    """
    Extract readable text from supported formats. ``filename`` is used for type detection only.
    Raises ``DocumentExtractionError`` on unsupported type or read failure.
    """
    if not data:
        raise DocumentExtractionError("File is empty.")
    name = (filename or "").strip().lower()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith(".doc"):
        raise DocumentExtractionError(
            "Legacy Word .doc is not supported. Save the document as .docx or PDF and upload again."
        )
    if name.endswith((".txt", ".md", ".markdown", ".text")):
        return _extract_plain_text(data)
    # Default: try plain text
    return _extract_plain_text(data)


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise DocumentExtractionError(
            "PDF support requires the 'pypdf' package. Install dependencies from requirements.txt."
        ) from e
    try:
        reader = PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
        return "\n\n".join(parts)
    except Exception as e:
        raise DocumentExtractionError(f"Could not read PDF: {e}") from e


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as e:
        raise DocumentExtractionError(
            "Word support requires 'python-docx'. Install dependencies from requirements.txt."
        ) from e
    if not zipfile.is_zipfile(io.BytesIO(data)):
        raise DocumentExtractionError("File does not look like a valid .docx (ZIP) archive.")
    try:
        document = docx.Document(io.BytesIO(data))
        return _docx_all_text(document)
    except Exception as e:
        raise DocumentExtractionError(f"Could not read Word document: {e}") from e


def _docx_all_text(document) -> str:
    """Paragraphs and table cells as lines."""
    parts: list[str] = []
    for p in document.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in document.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(x for x in cells if x)
            if line:
                parts.append(line)
    return "\n".join(parts)


def _extract_plain_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


# --- Preprocessing ---

# Leading annotation labels (longer phrases first). Strips label only; keeps the rest of the line.
_LEADING_ANNOTATION_STRIPS: Final[tuple[str, ...]] = (
    r"(?i)^\s*TO\s+BE\s+DISCUSSED\s*:?\s*",
    r"(?i)^\s*IMPORTANT\s*:\s*",
    r"(?i)^\s*NOTE\s*:\s*",
    r"(?i)^\s*TBD\s*:\s*",
    r"(?i)^\s*TODO\s*:\s*",
    r"(?i)^\s*FIXME\s*:\s*",
)

# Line is only a label token with no requirement text — drop the line.
# ``TBD`` alone is **not** here so it can be classified into ``open_items``.
_LABEL_ONLY_LINE: Final[re.Pattern[str]] = re.compile(
    r"(?i)^\s*(NOTE|IMPORTANT|TODO|FIXME|TO\s+BE\s+DISCUSSED)\s*:?\s*$"
)

# Lines that are coordination / uncertainty / placeholders — not functional requirements.
_OPEN_ITEM_LINE: Final[tuple[re.Pattern[str], ...]] = (
    re.compile(r"(?i)^\s*(tbd|tbc|tba|n/?a|wip|pending|unknown)\s*[.:]?\s*$"),
    re.compile(
        r"(?i)^\s*to\s+be\s+(decided|determined|defined|confirmed|discussed|finalized)\s*[.:]?\s*$"
    ),
    re.compile(
        r"(?i)^\s*(discuss|review|align|confirm|verify|validate)\s+.+\bwith\s+(the\s+)?(client|team|stakeholders?|stakeholder|product\s+owner|po|business|vendor|pm|sales)\b"
    ),
    re.compile(
        r"(?i)^\s*(needs?\s+clarification|clarification\s+needed|open\s+question|outstanding\s+question)\b"
    ),
    re.compile(r"(?i)^\s*(unclear|uncertain|ambiguous)\s*[.:]?\s*$"),
    re.compile(
        r"(?i)^\s*(unclear|uncertain|ambiguous)\s+(about|whether|if|how|what|which)\b"
    ),
    re.compile(r"(?i)^\s*(decide|determine)\s+(whether|if|how|what)\b"),
    re.compile(r"(?i)^\s*(flag|follow\s+up|follow-up|escalate)\b"),
    re.compile(
        r"(?i)^\s*(need\s+to\s+clarify|clarify\s+whether|unclear\s+whether|uncertain\s+whether)\b"
    ),
)


@dataclass(frozen=True)
class PreprocessedRequirement:
    """Normalized text split into functional requirements vs. open discussion / clarification lines."""

    functional_text: str
    open_items: tuple[str, ...] = ()


def _line_is_open_item(s: str) -> bool:
    t = (s or "").strip()
    if not t:
        return False
    return any(p.search(t) for p in _OPEN_ITEM_LINE)


def preprocess_requirement_with_classification(text: str) -> PreprocessedRequirement:
    """
    Same normalization as :func:`preprocess_requirement_text`, but **classifies** each substantive
    line as either a **functional** requirement statement or an **open item** (pending discussion,
    clarification, uncertainty). Open items are **excluded** from ``functional_text`` and intended
    for clarification / gap analysis routing only.
    """
    if not text:
        return PreprocessedRequirement(functional_text="", open_items=())
    raw = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = raw.split("\n")
    functional_lines: list[str] = []
    open_lines: list[str] = []
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if _LABEL_ONLY_LINE.match(s):
            continue
        s = _strip_annotation_labels(s)
        if not s.strip():
            # ``TBD`` / ``TBD:`` with no other text → route to open items (do not drop silently).
            if re.match(r"(?i)^\s*tbd\s*:?\s*$", line.strip()):
                open_lines.append("TBD")
            continue
        s = _normalize_bullet_line(s)
        s = re.sub(r"[ \t]+", " ", s).strip()
        if _line_is_open_item(s):
            open_lines.append(s)
        else:
            functional_lines.append(s)

    def _merge(blocks: list[str]) -> str:
        merged = "\n".join(blocks)
        merged = re.sub(r"\n{3,}", "\n\n", merged)
        merged = re.sub(r"[ \t]+", " ", merged)
        return merged.strip()

    return PreprocessedRequirement(
        functional_text=_merge(functional_lines),
        open_items=tuple(open_lines),
    )


def preprocess_requirement_text(text: str) -> str:
    """
    Normalize extracted or pasted text: whitespace, bullets, strip annotation **labels** while
    **keeping** sentence content, and **exclude** lines classified as open discussion / clarification
    from the returned string (those are available via :func:`preprocess_requirement_with_classification`).
    """
    return preprocess_requirement_with_classification(text).functional_text


def _strip_annotation_labels(s: str) -> str:
    """
    Remove common note/annotation prefixes repeatedly; preserve all substantive text after them.
    """
    t = s.strip()
    if not t:
        return ""
    for _ in range(12):
        prev = t
        for pat in _LEADING_ANNOTATION_STRIPS:
            t2 = re.sub(pat, "", t, count=1)
            if t2 != t:
                t = t2.strip()
                break
        if t == prev:
            break
    return t


def _normalize_bullet_line(s: str) -> str:
    """Turn common bullet prefixes into a readable '- item' form."""
    m = re.match(r"^[\s•\-\*·◦▪]+(.+)$", s)
    if m:
        return "- " + m.group(1).strip()
    return s
